#!/usr/bin/env python3
"""Build the WITNESS technical monograph and its publication figures.

The Markdown, DOCX, and figures are generated from the same structured source
so the GitHub and downloadable editions cannot silently drift apart.
"""

from __future__ import annotations

import math
from pathlib import Path
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
FIGURES = DOCS / "figures"
DOCX_PATH = DOCS / "WITNESS-Technical-Monograph.docx"
MARKDOWN_PATH = DOCS / "WITNESS-Technical-Monograph.md"

INK = "10211C"
MUTED = "53665E"
GREEN = "08784B"
CYAN = "16707A"
AMBER = "8A5B00"
PALE = "EDF6F2"
LINE = "AFC0B9"
LIGHT = "F7FAF8"
BLACK = "07100C"

# Design token map: narrative_proposal base with one named
# `technical_monograph` override for restrained WITNESS publication colors.
TOKENS = {
    "page": {"margin": 1.0, "header": 0.492, "footer": 0.492},
    "body": {"font": "Arial", "size": 10.5, "after": 7, "line": 1.22},
    "title": {"font": "Arial", "size": 31, "color": INK},
    "h1": {"font": "Arial", "size": 17, "color": GREEN, "before": 18, "after": 9},
    "h2": {"font": "Arial", "size": 13, "color": CYAN, "before": 13, "after": 6},
    "h3": {"font": "Arial", "size": 11.5, "color": INK, "before": 9, "after": 4},
    "table": {"width_dxa": 9360, "indent_dxa": 120, "cell_margin": 120},
}


CAPSULES = [
    {
        "index": "01", "id": "systems", "case": "CX-001", "label": "Distributed Systems",
        "system": "Raft + MVCC", "runtime": "C++17 to WebAssembly",
        "repo": "asp53826/raft-mvcc", "commit": "a5dd0986756b26ff5f375b73d03d618d9b2aded0",
        "receipt": "57c9cc96b82129672391606b85b45a8b3cdd9f32588a3df22bb841fd78890313",
        "claim": "A command appended by a minority leader cannot become committed or visible without majority acknowledgement.",
        "attack": "Elect node 1, isolate it from four peers, and submit transaction 7 to the stale leader.",
        "oracle": "The isolated log grows locally, but commit index and every MVCC state digest remain unchanged.",
        "observed": "Local log 1 to 2; commit index 1 to 1; four messages dropped; visible value: none.",
        "boundary": "The deterministic transport models exact partitions. It does not reproduce TCP backoff, kernel scheduling, disk faults, or fsync latency.",
        "figure": "fig04-raft-partition.png",
        "analysis": [
            "The distinction between acceptance and commitment is the center of this capsule. A client-facing append can look successful even though the replicated state machine has not crossed the majority boundary. WITNESS therefore observes the private log, commit index, follower logs, and state visibility separately instead of collapsing them into one success flag.",
            "The failure injection is intentionally minimal: one bidirectional partition and one proposal. That narrow trace makes the oracle inspectable. The result is evidence for the majority-before-commit invariant at the pinned revision, not evidence that every Raft implementation or every storage layer is correct.",
        ],
    },
    {
        "index": "02", "id": "ml", "case": "CX-009", "label": "ML Compiler",
        "system": "TensorForge", "runtime": "TypeScript",
        "repo": "asp53826/tensorforge-webgpu", "commit": "70a5c855ef35e01a29ba351eb2645670060b9f0b",
        "receipt": "37576d77588845f3023d2debcb397f53eacfb91cdca94250f1013c1f778eefa6",
        "claim": "Matrix multiplication must reject unequal contracting dimensions before optimization or GPU dispatch.",
        "attack": "Mutate W1 from [128,256] to [127,256] while retaining the upstream [8,128] activation.",
        "oracle": "The validator emits the exact shape error and the GPU dispatch counter remains zero.",
        "observed": "Input [8,128]; W1 [128,256] to [127,256]; GPU dispatches 0; rejection exact.",
        "boundary": "This is a static rank-and-dimension guard. It does not establish numerical stability, WebGPU driver correctness, or performance portability.",
        "figure": "fig05-shape-gate.png",
        "analysis": [
            "Compiler correctness begins before optimization. If malformed tensor programs reach fusion or code generation, later passes can transform the symptom and make the root cause harder to identify. The capsule makes validation a front-door invariant and records that no GPU work was launched.",
            "The oracle is stronger than checking for any exception. It binds the expected mismatch to the contracting dimension and separately observes dispatch count. This prevents an unrelated runtime failure from being misreported as successful shape verification.",
        ],
    },
    {
        "index": "03", "id": "autonomy", "case": "CX-005", "label": "Autonomy",
        "system": "Track Fusion", "runtime": "Python, NumPy, and SciPy in a Pyodide Worker",
        "repo": "asp53826/track-fusion", "commit": "9a1092b8da88f5904ca190a393d77bd2bcd111b8",
        "receipt": "762231d01fdd53b48b8b958f22cc06b03476b5f3997246a6aa84168eaf29038d",
        "claim": "Target-free Poisson clutter must not be promoted into persistent confirmed tracks.",
        "attack": "Replay 40 seeded scans with zero targets and an expected six uniform false alarms per scan.",
        "oracle": "Wald sequential log-likelihood promotion produces no more than four confirmations.",
        "observed": "40 scans; 264 clutter returns; 0 confirmed tracks; acceptance boundary <= 4.",
        "boundary": "The scenario uses simulated planar measurements and a declared clutter model, not calibrated production radar or EO sensor data.",
        "figure": "fig06-clutter-gate.png",
        "analysis": [
            "A large validation gate can collect enough clutter to resemble persistence unless confirmation evidence accounts for both target likelihood and false-alarm density. The capsule attacks this failure directly by removing real targets from the scene.",
            "Zero confirmations is not presented as a universal false-track rate. It is one seeded observation under a published target-free model. The receipt preserves the seed, scan count, clutter density, and acceptance threshold so the evidence can be replayed without being generalized beyond its regime.",
        ],
    },
    {
        "index": "04", "id": "quant", "case": "CX-008", "label": "Market Microstructure",
        "system": "Limit Order Book Market Making", "runtime": "Python in a Pyodide Worker",
        "repo": "asp53826/lob-market-making", "commit": "82c75a1fd7824dd96bdb92432b58b3557cb7d577",
        "receipt": "fa02f77611c71febb1d064c94e534db375104fd8998a39e5aee6d78d0d6e0cba",
        "claim": "A profitability report must keep carried inventory risk and exact P&L decomposition beside total profit.",
        "attack": "Run paired 4,000-step simulations at seed 3 with naive and inventory-skew quoting.",
        "oracle": "Inventory skew lowers inventory dispersion and both cash/inventory P&L ledgers reconcile exactly.",
        "observed": "Inventory sigma 33.88 to 11.36; 66.5% reduction; two of two P&L ledgers exact.",
        "boundary": "This is a synthetic paired test, not a live trading result, execution recommendation, or claim about transaction costs in a real venue.",
        "figure": "fig07-risk-ledger.png",
        "analysis": [
            "Profit alone is an incomplete state variable. Two strategies can end with similar totals while carrying very different exposure during the run. WITNESS keeps the dispersion of inventory and the accounting identity visible so a favorable total cannot hide a fragile path.",
            "The comparison uses common configuration and seed to reduce irrelevant sampling differences. It still remains a targeted simulator result. No conclusion is drawn about expected live returns, optimal parameter calibration, or market impact.",
        ],
    },
    {
        "index": "05", "id": "search", "case": "CX-010", "label": "Vector Search",
        "system": "ANNLite HNSW", "runtime": "C++17 to WebAssembly",
        "repo": "asp53826/annlite", "commit": "dc922c8c1816df48d75f2a16e51cd25819aab070",
        "receipt": "cffc4aa733f3c13b314b666c99304a9fea583558322ed829e6aa83a7f118bf4a",
        "claim": "Approximate-search speed must be reported with exact-ground-truth recall across the search-work frontier.",
        "attack": "Hold one deterministic 1,800-vector graph and 80 queries fixed while sweeping efSearch from 10 to 100.",
        "oracle": "Recall and distance computations must be monotone, and the high-budget endpoint must recover the low-budget misses.",
        "observed": "Recall 95.375% to 100.000%; mean distance computations 132.97 to 207.91; frontier monotone.",
        "boundary": "This browser-scale graph is not the repository's published SIFT1M throughput benchmark and does not claim production latency.",
        "figure": "fig08-ann-frontier.png",
        "analysis": [
            "HNSW exposes a controllable accuracy-cost trade. Reporting only latency creates an incentive to choose a search breadth that looks fast while silently omitting neighbors. The browser capsule computes exact top-10 neighbors first, then measures the approximate result against that oracle.",
            "The observed endpoint reaches full recall on this targeted dataset. That statement is deliberately local: different distributions, graph parameters, dimensions, and insertion orders can change the frontier. The contribution is the executable disclosure pattern, not a claim that efSearch=100 is universally sufficient.",
        ],
    },
    {
        "index": "06", "id": "imaging", "case": "CX-011", "label": "SAR Imaging",
        "system": "Phase Gradient Autofocus", "runtime": "Python and NumPy in a Pyodide Worker",
        "repo": "asp53826/sar-focus", "commit": "83041573ebcfa935640f91cb091e1dbaad4f9aaa",
        "receipt": "0025d98e9f631c2f487bf3214364a9af6b5987817449e11671fda46404dd9b9e",
        "claim": "Autofocus acceptance must report both image entropy and impulse-response width.",
        "attack": "Inject a seeded random phase error at 5 rad RMS, backproject a four-target scene, and run eight PGA iterations.",
        "oracle": "Detect the counterexample in which entropy decreases while azimuth IRW widens.",
        "observed": "Entropy 10.103 to 9.917; azimuth IRW 0.3096 m to 0.3461 m; recovery 0.89x.",
        "boundary": "The input is a simulated stripmap point-target scene, not flight data or a claim of operational image quality.",
        "figure": "fig09-sar-metrics.png",
        "analysis": [
            "Autofocus can optimize one scalar while damaging a property that the scalar does not encode. Here the entropy score improves, yet the measured main-lobe width becomes worse. A single-number acceptance test would approve the wrong output.",
            "The capsule treats disagreement as the expected finding. Passing means the system correctly detects the losing regime; it does not mean PGA failed in general. The repository also contains regimes where the same method recovers a known phase error near the diffraction limit.",
        ],
    },
    {
        "index": "07", "id": "navigation", "case": "CX-006", "label": "Visual-Inertial Navigation",
        "system": "MSCKF VIO", "runtime": "Python and NumPy in a Pyodide Worker",
        "repo": "asp53826/vio-nav", "commit": "f38714f391ffb0dc3ae06b418aae51f929000e1e",
        "receipt": "be84bcd3f1737f84f292a0169384b6701f537613a847621b19bf55d887a9f37e",
        "claim": "A stationary camera must reject depth updates that contain no translational parallax.",
        "attack": "Hold the camera in an eight-second synthetic hover with 250 seeded landmarks.",
        "oracle": "Reject non-triangulable feature tracks and produce no position correction beyond dead reckoning.",
        "observed": "6,956 parallax rejections; dead-reckoning ATE 0.01687 m; VIO ATE 0.01687 m.",
        "boundary": "The result demonstrates one observability guard in simulation; it is not evidence of field navigation accuracy or full estimator consistency.",
        "figure": "fig10-vio-observability.png",
        "analysis": [
            "Multiple image observations do not automatically create geometric information. With rotation but no translational baseline, bearing rays do not supply the parallax required for stable depth triangulation. Updating anyway would manufacture structure from an unobservable direction.",
            "The equal ATE values are therefore a feature of this test: the camera provides no valid translational correction. The stronger claim is that the filter recognizes the missing information and refuses to invent it.",
        ],
    },
    {
        "index": "08", "id": "numerical", "case": "CX-007", "label": "Numerical Differentiation",
        "system": "AAD Greeks", "runtime": "Python and NumPy in a Pyodide Worker",
        "repo": "asp53826/aad-greeks", "commit": "c26ef7aa18f473b7cfd85642a62aef3a17b94c58",
        "receipt": "d2d00e9583dfd3f7e7674985b8955cda352818e4a068299749e95cb97e75e4a3",
        "claim": "A discontinuous payoff cannot rely on pathwise automatic differentiation without an independent oracle.",
        "attack": "Differentiate a 200,000-path digital payoff at the strike using identical seeded normal draws.",
        "oracle": "Compare sharp pathwise AAD with analytic delta, common-random-number central difference, and smoothed AAD.",
        "observed": "Pathwise AAD 0.000000; analytic 0.019333; central difference 0.019220; smoothed AAD 0.019479.",
        "boundary": "This is a numerical counterexample for one digital payoff, not financial advice or validation of a production pricing library.",
        "figure": "fig11-aad-discontinuity.png",
        "analysis": [
            "The program derivative is exact for almost every simulated path: the derivative of the step payoff is zero away from the strike. The estimator is nevertheless wrong for the derivative of the expectation because differentiation and expectation cannot be interchanged naively at the discontinuity.",
            "The important behavior is silent confidence. No NaN or exception announces the failure. WITNESS therefore makes three independent estimates visible and accepts the capsule only when the false zero and both recovery methods appear together.",
        ],
    },
]


REFERENCES = [
    "Association for Computing Machinery. Artifact Review and Badging. https://www.acm.org/publications/policies/artifact-review-and-badging-current",
    "Ongaro, D., and Ousterhout, J. In Search of an Understandable Consensus Algorithm (Extended Version). 2014. https://web.stanford.edu/~ouster/cgi-bin/papers/raft-extended.pdf",
    "Herlihy, M. P., and Wing, J. M. Linearizability: A Correctness Condition for Concurrent Objects. ACM TOPLAS 12(3), 1990. https://doi.org/10.1145/78969.78972",
    "Malkov, Y. A., and Yashunin, D. A. Efficient and Robust Approximate Nearest Neighbor Search Using Hierarchical Navigable Small World Graphs. IEEE TPAMI 42(4), 2020. https://arxiv.org/abs/1603.09320",
    "Mourikis, A. I., and Roumeliotis, S. I. A Multi-State Constraint Kalman Filter for Vision-aided Inertial Navigation. ICRA, 2007. https://www-users.cse.umn.edu/~stergios/papers/ICRA07-MSCKF.pdf",
    "Wahl, D. E., Eichel, P. H., Ghiglia, D. C., and Jakowatz, C. V. Phase Gradient Autofocus - A Robust Tool for High Resolution SAR Phase Correction. IEEE TAES 30(3), 1994. https://doi.org/10.1109/7.303752",
    "Fortmann, T. E., Bar-Shalom, Y., and Scheffe, M. Sonar Tracking of Multiple Targets Using Joint Probabilistic Data Association. IEEE Journal of Oceanic Engineering 8(3), 1983. https://doi.org/10.1109/JOE.1983.1145560",
    "Avellaneda, M., and Stoikov, S. High-frequency Trading in a Limit Order Book. Quantitative Finance 8(3), 2008. https://doi.org/10.1080/14697680701381228",
    "Giles, M. B. Fifteen Years of Adjoint Algorithmic Differentiation in Finance. 2024. https://www2.maths.ox.ac.uk/~gilesm/files/AAD_Review.pdf",
    "LLVM Project. MLIR Shape Inference. https://mlir.llvm.org/docs/ShapeInference/",
    "Emscripten Project. WebAssembly and Modularized Output Documentation. https://emscripten.org/docs/compiling/WebAssembly.html",
    "Pyodide Project. Using Pyodide in a Web Worker, version 314.0.2. https://pyodide.org/en/stable/usage/webworker.html",
    "GitHub. Using Artifact Attestations to Establish Provenance for Builds. https://docs.github.com/en/actions/how-tos/secure-your-work/use-artifact-attestations/use-artifact-attestations",
    "Patel, A. WITNESS source repository, pinned capsule definitions, tests, and source manifest. https://github.com/asp53826/witness",
    "Patel, A. COUNTEREXAMPLE v1.0.0 release and canonical failure receipts. https://github.com/asp53826/counterexample/releases/tag/v1.0.0",
]


FRONT_MATTER = [
    ("Abstract", [
        "WITNESS is a browser-native reproducibility chamber for engineering claims. Instead of presenting screenshots or aggregate claims, it verifies a canonical failure receipt, binds that receipt to an immutable source revision, executes a bounded adversarial scenario, evaluates an explicit oracle, and emits a local run receipt whose attestation boundary is stated rather than implied.",
        "This monograph develops the evidence model behind WITNESS and analyzes eight executable capsules spanning distributed systems, compiler validation, multi-target tracking, market microstructure, approximate nearest-neighbor search, synthetic-aperture radar autofocus, visual-inertial navigation, and algorithmic differentiation. The contribution is not a new theorem shared by all eight domains. It is an engineering method for keeping claims, losing regimes, measurements, source identity, and limitations coupled closely enough that a reader can rerun and challenge them.",
        "All numerical observations reported here were reproduced in the browser against the pinned WITNESS build described in the artifact manifest. They are targeted counterexamples, not substitutes for full benchmark suites or independent external replication.",
    ]),
    ("Artifact Status and Claim Boundary", [
        "This is an independent technical monograph, not a university dissertation, peer-reviewed publication, or claim that an academic degree has been awarded. The level of treatment is deliberately research-oriented: formal predicates, adversarial methodology, exact source revisions, negative results, and reproducibility artifacts are included so the work can be evaluated on evidence rather than labels.",
        "The canonical COUNTEREXAMPLE release is provenance-attested. A WITNESS browser run is deterministic and hash-bound but explicitly local-unattested. No ACM badge or independent reproduction status is claimed. Section 15 maps the current artifacts to evaluation criteria without awarding them to itself.",
    ]),
]


CORE_SECTIONS = [
    ("1. Research Problem and Contribution", [
        "Engineering portfolios often collapse several distinct propositions into one visual claim: that code exists, that it ran, that a reported number came from that code, that the number supports the stated conclusion, and that the conclusion remains valid outside the demonstrated regime. WITNESS separates those propositions and exposes the joins between them.",
        "The working research question is: how can a public engineering artifact make a narrow claim executable, falsifiable, source-bound, and limitation-aware without requiring a reviewer to reconstruct the author's entire environment? The design response is a small verification protocol whose output is useful even when the expected result is a failure.",
        "The first contribution is a receipt-gated execution model. Every capsule names a claim, adversarial condition, oracle, source commit, and canonical receipt digest. Runtime execution is blocked if format, receipt hash, embedded capsule hash, or commit binding fails.",
        "The second contribution is negative-result centricity. A capsule is not a product demo with an error state added later. It begins from a losing regime: minority acceptance without commitment, low-budget recall loss, clutter promotion, entropy/resolution disagreement, missing parallax, or a derivative that is exactly computed and still wrong.",
        "The third contribution is a browser-local review surface. C++ engines compile to WebAssembly; Python engines execute in a module Worker through Pyodide; TypeScript validation runs directly. This reduces installation friction while preserving the distinction between a focused browser scenario and the repository's larger native benchmark suite.",
    ]),
    ("2. Evidence Model", [
        "Let C be a narrow claim, A an adversarial scenario, O an oracle, S an immutable source revision, R a canonical receipt, and E an observed execution trace. WITNESS accepts a run only if both receipt integrity and behavioral evidence hold.",
        "The integrity predicate is defined as:",
        "I(R,S) = format(R) AND H(canonical(R without receiptSha256)) = R.receiptSha256 AND H(canonical(R.capsule)) = R.integrity.capsuleSha256 AND R.subject.commit = R.capsule.commit = S",
        "The behavioral predicate is capsule-specific:",
        "B(E,O) = O(E) = true",
        "The chamber verdict is therefore V = I(R,S) AND B(E,O). This decomposition matters: a correct behavioral trace from the wrong revision is not accepted, and a perfectly hashed receipt cannot rescue a failing oracle.",
        "The local run receipt records the selected capsule, subject commit, runtime, canonical receipt digest, integrity result, oracle result, timestamp, and a trace digest. It deliberately records signatureStatus=local-unattested. The receipt is evidence of what the visitor's browser observed, not an externally signed release artifact.",
    ]),
    ("3. Threat Model", [
        "WITNESS is designed against evidence failures that are common in public technical demonstrations: screenshots that cannot be replayed, moving main branches, results detached from commands, stale binaries, silently changed oracles, favorable metrics that hide a losing dimension, and deterministic claims built on unrecorded random state.",
        "The chamber does not defend against a malicious hosting origin, a compromised browser, fabricated upstream repositories, or a coordinated author who falsifies every layer consistently. It also does not transform self-reproduction into independent replication. Those limits are explicit because a threat model that claims everything protects nothing.",
        "The invert-verdict control is a small tamper experiment. It leaves engine output unchanged while negating the expected oracle comparison. The resulting failed trace demonstrates that WITNESS is not merely painting a success state after execution.",
    ]),
    ("4. Architecture and Runtime Isolation", [
        "The page has three trust-relevant layers. The integrity inspector fetches and verifies the canonical receipt. The execution adapter loads one of three bounded runtime families. The presentation layer renders measured metrics and trace lines from the returned structured result.",
        "C++ capsules use modularized Emscripten output and pre-fetched WebAssembly bytes. TensorForge uses the pinned validator directly in TypeScript. Five Python capsules run in one module Worker. The Worker mounts only enumerated pinned source files and selects only hard-coded scenarios; visitors cannot submit arbitrary Python.",
        "Worker isolation is a responsiveness boundary, not a security sandbox claim. It prevents scientific Python computations from blocking the main UI and keeps the DOM outside the worker global context. The hosting origin still controls both sides of the message channel.",
    ]),
    ("5. Experimental Method", [
        "Each capsule is constructed as a minimal discriminating experiment. One variable or regime is attacked, the random state is fixed where relevant, and the oracle observes the state necessary to distinguish the intended invariant from an unrelated failure.",
        "Results in this document are browser observations from the pinned scenarios, not copied headline numbers from repository READMEs. Where the browser intentionally uses a smaller dataset than a native benchmark, that difference is named in the capsule boundary.",
        "The method favors exact paired comparisons and identities over broad but weak scores: commit index beside local append, dispatch count beside a shape exception, exact brute-force neighbors beside approximate search, entropy beside impulse-response width, dead reckoning beside VIO in a no-parallax regime, and analytic/bumped/smoothed deltas beside pathwise AAD.",
    ]),
]


SYNTHESIS = [
    ("14. Cross-Capsule Synthesis", [
        "Across the eight domains, the same structural failure appears repeatedly: an available measurement is mistaken for the property that actually matters. Local append is mistaken for commitment. Low latency is mistaken for adequate recall. Lower entropy is mistaken for better resolution. A formally propagated derivative is mistaken for the derivative of an expectation. WITNESS counters these substitutions by making the missing companion variable part of the oracle.",
        "The capsules also show three distinct oracle classes. State invariants inspect forbidden transitions, as in Raft and TensorForge. Comparative oracles hold an environment fixed and compare paired methods, as in market making and AAD. Observability or disagreement oracles detect when the information required for a conclusion is absent or contradictory, as in VIO and SAR.",
        "No single runtime establishes credibility. The strongest signal is agreement across layers: human-readable claim, adversarial test, machine-checkable oracle, pinned source, canonical hash, deterministic execution, and an explicit boundary. Removing any one layer creates a predictable ambiguity.",
    ]),
    ("15. Artifact Evaluation Readiness", [
        "ACM artifact guidance distinguishes availability, functional evaluation, reusability, and independently reproduced results. WITNESS can prepare evidence for such review, but it cannot award itself a badge.",
        "Availability: the repositories, source revisions, receipts, and browser chamber are public. Functionality: the CI suite validates eight receipt bindings, unit tests, WebAssembly smoke scenarios, and the production build. Reusability: source manifests, pinned links, runtime boundaries, and one-command checks lower the cost of inspection. Independent reproduction: not yet established; that status requires another person or team to obtain the results.",
        "A credible next step is a blinded external replay packet: a fresh evaluator receives the release bundle and commands, records environment and deviations, reruns a predeclared subset, and publishes an independent report whether the results match or fail.",
    ]),
    ("16. Limitations", [
        "The browser scenarios are deliberately bounded and should not be confused with full system evaluation. WebAssembly changes the execution environment for the C++ engines. Pyodide uses browser-distributed scientific Python. Timing measurements are excluded because browser, device, thermal, and scheduling conditions are not controlled.",
        "Receipt hashing establishes identity and detects accidental or adversarial mutation relative to the expected digest. It does not prove that the original statement was true, that the source is free of bugs, or that the author did not choose a favorable scenario.",
        "Several domains depend on simulated data. Simulation is valuable because it exposes ground truth and enables deterministic attacks; it is insufficient for claims about deployment realism. Each capsule therefore limits its conclusion to the modeled regime.",
        "The monograph is authored by the same project owner and should be read as technical documentation, not peer review. Its value is that claims and methods are inspectable enough for others to disagree precisely.",
    ]),
    ("17. Future Work", [
        "Add environment capture to local run receipts: browser engine, WebAssembly feature set, Pyodide build, package versions, device class, and runtime checksum, while avoiding collection of personal identifiers.",
        "Introduce pre-registered scenario suites so a reviewer can run multiple seeds or parameter cells without granting arbitrary-code execution. Aggregate results should retain every cell and never replace the underlying traces.",
        "Publish signed release artifacts for the WITNESS production bundle and its source manifest. GitHub artifact attestations can bind build provenance, but the UI must still distinguish build provenance from independent result reproduction.",
        "Pursue one genuine external artifact review. The most valuable outcome is not a success badge; it is a public report that identifies which capsule could not be reproduced and why.",
    ]),
    ("18. Conclusion", [
        "WITNESS treats a claim as an executable object with an attack surface, not a sentence decorated by a benchmark. Its eight capsules demonstrate that sophisticated engineering work becomes more credible when the losing regime is easier to find, the oracle is narrower, the source revision is immutable, and the boundary is impossible to miss.",
        "The result is neither a proof of universal correctness nor a replacement for peer review. It is a practical research artifact: one that lets a reader move from assertion to counterexample, from counterexample to code, and from code to a measured verdict without leaving the evidence chain implicit.",
    ]),
]


def font_path(bold: bool = False, mono: bool = False) -> str:
    if mono:
        return "/System/Library/Fonts/Menlo.ttc"
    if bold:
        return "/System/Library/Fonts/Supplemental/Arial Bold.ttf"
    return "/System/Library/Fonts/Supplemental/Arial.ttf"


def pil_font(size: int, bold: bool = False, mono: bool = False) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(font_path(bold=bold, mono=mono), size=size)


def new_figure(title: str, subtitle: str = "") -> tuple[Image.Image, ImageDraw.ImageDraw]:
    image = Image.new("RGB", (1600, 900), "#07100C")
    draw = ImageDraw.Draw(image)
    for x in range(0, 1600, 80): draw.line((x, 0, x, 900), fill="#0E1C17", width=1)
    for y in range(0, 900, 80): draw.line((0, y, 1600, y), fill="#0E1C17", width=1)
    draw.rounded_rectangle((24, 24, 1576, 876), radius=20, outline="#293C34", width=2)
    draw.text((70, 58), title.upper(), font=pil_font(34, True), fill="#EDF6F2")
    if subtitle: draw.text((72, 108), subtitle, font=pil_font(19), fill="#8DA39A")
    draw.line((70, 148, 1530, 148), fill="#315046", width=2)
    return image, draw


def box(draw: ImageDraw.ImageDraw, xy, title: str, detail: str, accent="#8FFFC3", fill="#0D1914"):
    draw.rounded_rectangle(xy, radius=14, fill=fill, outline=accent, width=3)
    x1, y1, x2, y2 = xy
    draw.text((x1 + 24, y1 + 20), title, font=pil_font(23, True), fill=accent)
    lines = detail.split("\n")
    for i, line in enumerate(lines):
        draw.text((x1 + 24, y1 + 63 + i * 29), line, font=pil_font(18), fill="#C7D5CF")


def arrow(draw: ImageDraw.ImageDraw, start, end, color="#68D7DC", width=5):
    draw.line((*start, *end), fill=color, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    for delta in (2.55, -2.55):
        p = (end[0] + 18 * math.cos(angle + delta), end[1] + 18 * math.sin(angle + delta))
        draw.line((*end, *p), fill=color, width=width)


def label(draw, xy, text, color="#8DA39A", size=18, bold=False):
    draw.text(xy, text, font=pil_font(size, bold), fill=color)


def generate_figures() -> None:
    FIGURES.mkdir(parents=True, exist_ok=True)

    im, d = new_figure("The WITNESS evidence chain", "Identity and behavior must both hold before a verdict is accepted")
    labels = [("CLAIM", "narrow statement"), ("ATTACK", "losing regime"), ("ORACLE", "deciding predicate"), ("SOURCE", "pinned commit"), ("RECEIPT", "canonical hashes"), ("RUN", "measured trace")]
    for i, (a, b) in enumerate(labels):
        x = 55 + i * 255
        box(d, (x, 310, x + 205, 500), a, b, "#8FFFC3" if i in (0, 5) else "#68D7DC")
        if i < len(labels)-1: arrow(d, (x + 205, 405), (x + 247, 405))
    box(d, (460, 630, 1140, 770), "ACCEPT", "receipt integrity  AND  capsule oracle", "#FFC96A")
    arrow(d, (800, 510), (800, 630), "#FFC96A")
    im.save(FIGURES / "fig01-evidence-chain.png", quality=95)

    im, d = new_figure("Browser execution topology", "Three bounded runtime families; one receipt-gated result contract")
    box(d, (70, 260, 410, 500), "INTEGRITY", "format\nreceipt SHA-256\ncapsule SHA-256\ncommit binding", "#AA9CFF")
    box(d, (630, 205, 970, 385), "C++17 / WASM", "Raft + MVCC\nANNLite HNSW", "#8FFFC3")
    box(d, (630, 420, 970, 600), "TYPESCRIPT", "TensorForge\nshape validator", "#68D7DC")
    box(d, (630, 635, 970, 815), "PYODIDE WORKER", "five pinned Python\nscenarios", "#FFC96A")
    box(d, (1190, 310, 1530, 580), "RESULT", "metrics\ntrace\noracle verdict\nlocal receipt", "#8FFFC3")
    for y in (295, 510, 725): arrow(d, (970, y), (1190, 430), "#68D7DC", 4)
    for y in (295, 510, 725): arrow(d, (410, 380), (630, y), "#AA9CFF", 4)
    im.save(FIGURES / "fig02-runtime-topology.png", quality=95)

    im, d = new_figure("Eight executable capsules", "One evidence contract across unrelated engineering domains")
    matrix_labels = (
        "DISTRIBUTED SYSTEMS",
        "ML COMPILER",
        "AUTONOMY",
        "MARKET SIMULATION",
        "VECTOR SEARCH",
        "SAR IMAGING",
        "VIO",
        "NUMERICAL AAD",
    )
    matrix_runtimes = (
        "C++17 / WASM",
        "TypeScript",
        "Python / Pyodide",
        "Python / Pyodide",
        "C++17 / WASM",
        "Python / Pyodide",
        "Python / Pyodide",
        "Python / Pyodide",
    )
    for i, c in enumerate(CAPSULES):
        col, row = i % 4, i // 4
        x, y = 60 + col * 380, 220 + row * 275
        box(d, (x, y, x + 340, y + 210), f"{c['index']}  {matrix_labels[i]}", f"{c['case']}\n{c['system']}\n{matrix_runtimes[i]}", "#8FFFC3" if row else "#68D7DC")
    im.save(FIGURES / "fig03-capsule-matrix.png", quality=95)

    im, d = new_figure("Raft minority partition", "A local append is observable; commitment and state visibility are forbidden")
    for i in range(5):
        x = 150 + i * 290
        color = "#FF765F" if i == 0 else "#8FFFC3"
        box(d, (x, 300, x + 200, 470), f"NODE {i+1}", "leader" if i == 0 else "follower", color)
    d.line((410, 220, 410, 610), fill="#FF765F", width=8)
    label(d, (330, 180), "PARTITION", "#FF765F", 20, True)
    arrow(d, (150, 610), (350, 610), "#FFC96A")
    label(d, (90, 650), "tx7 appended locally", "#FFC96A", 20, True)
    label(d, (720, 610), "commit index unchanged  |  MVCC value invisible", "#8FFFC3", 24, True)
    im.save(FIGURES / "fig04-raft-partition.png", quality=95)

    im, d = new_figure("Tensor shape admission gate", "Malformed matrix multiplication stops before optimization and dispatch")
    box(d, (90, 310, 400, 500), "ACTIVATION", "A : [8, 128]", "#68D7DC")
    box(d, (520, 310, 830, 500), "MUTATED WEIGHT", "W1 : [127, 256]", "#FF765F")
    box(d, (950, 250, 1260, 560), "VALIDATOR", "contracting dims\n128 != 127\n\nREJECT", "#FFC96A")
    box(d, (1350, 310, 1530, 500), "GPU", "dispatches\n0", "#8FFFC3")
    arrow(d, (400, 405), (520, 405)); arrow(d, (830, 405), (950, 405));
    d.line((1260, 405, 1350, 405), fill="#FF765F", width=8)
    im.save(FIGURES / "fig05-shape-gate.png", quality=95)

    im, d = new_figure("Clutter confirmation guard", "Target-free returns may enter gates but must not accumulate target evidence")
    for i in range(36):
        x = 90 + (i * 97) % 1380; y = 210 + (i * 67) % 430
        d.ellipse((x, y, x+10, y+10), fill="#68D7DC")
    d.ellipse((540, 280, 1050, 700), outline="#FFC96A", width=5)
    box(d, (1100, 270, 1500, 560), "WALD LLR", "40 scans\n264 clutter returns\n0 confirmations\nboundary <= 4", "#8FFFC3")
    arrow(d, (1020, 490), (1100, 415), "#FFC96A")
    im.save(FIGURES / "fig06-clutter-gate.png", quality=95)

    im, d = new_figure("Profit and inventory must travel together", "Paired seed; exact accounting identity; risk reported beside total")
    base = 700
    d.rectangle((220, base-390, 480, base), fill="#FF765F")
    d.rectangle((610, base-131, 870, base), fill="#8FFFC3")
    label(d, (240, 730), "NAIVE", "#C7D5CF", 24, True); label(d, (635, 730), "INVENTORY SKEW", "#C7D5CF", 24, True)
    label(d, (265, 270), "sigma = 33.88", "#FF765F", 25, True); label(d, (650, 520), "sigma = 11.36", "#8FFFC3", 25, True)
    box(d, (1030, 280, 1510, 610), "P&L LEDGER", "cash change\n+ inventory value\n= total P&L\n\n2 / 2 exact", "#68D7DC")
    im.save(FIGURES / "fig07-risk-ledger.png", quality=95)

    im, d = new_figure("ANN recall-work frontier", "Exact top-10 ground truth; one deterministic graph and query set")
    pts = [(10,.95375,132.97),(20,.995,164.06),(50,1.0,196.69),(100,1.0,207.91)]
    x0,y0,x1,y1 = 150,700,1450,220
    d.line((x0,y0,x1,y0), fill="#8DA39A", width=3); d.line((x0,y0,x0,y1), fill="#8DA39A", width=3)
    for ef, recall, work in pts:
        x=x0+(ef/100)*(x1-x0); y=y0-((recall-.94)/.06)*(y0-y1)
        d.ellipse((x-12,y-12,x+12,y+12), fill="#8FFFC3")
        label(d,(x-55,y-65),f"{recall*100:.3f}%", "#EDF6F2", 18, True)
        label(d,(x-55,y+25),f"{work:.1f} comps", "#8DA39A", 16)
        label(d,(x-10,y0+25),str(ef), "#8DA39A", 16)
    d.line([(x0+(ef/100)*(x1-x0), y0-((rec-.94)/.06)*(y0-y1)) for ef,rec,_ in pts], fill="#68D7DC", width=5)
    label(d,(650,805),"efSearch", "#8DA39A", 20, True); label(d,(35,420),"recall", "#8DA39A", 20, True)
    im.save(FIGURES / "fig08-ann-frontier.png", quality=95)

    im, d = new_figure("SAR autofocus metric disagreement", "Passing the capsule means detecting that the two quality measures disagree")
    box(d,(90,240,710,690),"ENTROPY  (LOWER)","10.103  ->  9.917\n\nIMPROVED", "#8FFFC3")
    box(d,(890,240,1510,690),"AZIMUTH IRW  (LOWER)","0.3096 m  ->  0.3461 m\n\nWORSENED", "#FF765F")
    arrow(d,(710,465),(890,465),"#FFC96A",7)
    label(d,(650,750),"ONE SCORE CANNOT REPRESENT IMAGE QUALITY", "#FFC96A",22,True)
    im.save(FIGURES / "fig09-sar-metrics.png", quality=95)

    im, d = new_figure("VIO observability under hover", "Repeated bearings without translational baseline do not determine feature depth")
    d.ellipse((230,370,350,490), fill="#68D7DC")
    label(d,(210,510),"camera pose", "#C7D5CF",20,True)
    for y in (240,330,570,660):
        d.line((350,430,1300,y), fill="#53665E", width=3)
        d.ellipse((1290,y-10,1310,y+10), fill="#8DA39A")
    d.ellipse((214,354,366,506), outline="#FFC96A", width=4)
    label(d,(710,270),"same optical center", "#FFC96A",23,True)
    box(d,(1030,650,1510,820),"ORACLE","6,956 rejected\nVIO ATE = DR ATE", "#8FFFC3")
    im.save(FIGURES / "fig10-vio-observability.png", quality=95)

    im, d = new_figure("Digital payoff: exact path derivative, wrong sensitivity", "The step is flat almost everywhere while the expectation changes with spot")
    x0,y0,x1,y1=160,690,1420,240
    d.line((x0,y0,x1,y0),fill="#8DA39A",width=3); d.line((x0,y0,x0,y1),fill="#8DA39A",width=3)
    strike=790
    d.line((x0,y0,strike,y0),fill="#68D7DC",width=7); d.line((strike,y0,strike,y1),fill="#FF765F",width=7); d.line((strike,y1,x1,y1),fill="#68D7DC",width=7)
    d.line((strike,190,strike,740),fill="#FFC96A",width=2)
    label(d,(760,755),"K = 100", "#FFC96A",20,True)
    box(d,(210,260,650,480),"PATHWISE AAD","0.000000", "#FF765F")
    box(d,(950,420,1450,650),"INDEPENDENT ORACLES","analytic  0.019333\nbumped   0.019220\nsmoothed 0.019479", "#8FFFC3")
    im.save(FIGURES / "fig11-aad-discontinuity.png", quality=95)


def set_run(run, *, size=None, bold=None, italic=None, color=None, font="Arial"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    if size is not None: run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic
    if color: run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr(); tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None: tc_mar = OxmlElement("w:tcMar"); tc_pr.append(tc_mar)
    for m, v in (("top",top),("start",start),("bottom",bottom),("end",end)):
        node=tc_mar.find(qn(f"w:{m}"))
        if node is None: node=OxmlElement(f"w:{m}"); tc_mar.append(node)
        node.set(qn("w:w"),str(v)); node.set(qn("w:type"),"dxa")


def set_table_geometry(table, widths: list[int]):
    table.autofit = False; table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
    if tbl_w.getparent() is None: tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths))); tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd"); tbl_ind.set(qn("w:w"), "120"); tbl_ind.set(qn("w:type"), "dxa"); tbl_pr.append(tbl_ind)
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for width in widths:
        col=OxmlElement("w:gridCol"); col.set(qn("w:w"),str(width)); grid.append(col)
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width/1440)
            tc_w=cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"),str(width)); tc_w.set(qn("w:type"),"dxa")
            cell.vertical_alignment=WD_ALIGN_VERTICAL.CENTER; set_cell_margins(cell)
    if table.rows:
        repeat_table_header(table.rows[0])


def add_page_field(paragraph):
    run=paragraph.add_run(); fld=OxmlElement("w:fldSimple"); fld.set(qn("w:instr"),"PAGE"); run._r.addnext(fld)


def add_hyperlink(paragraph, text, url):
    part=paragraph.part; rel=part.relate_to(url,"http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",is_external=True)
    hyperlink=OxmlElement("w:hyperlink"); hyperlink.set(qn("r:id"),rel)
    run=OxmlElement("w:r"); rpr=OxmlElement("w:rPr"); color=OxmlElement("w:color"); color.set(qn("w:val"),CYAN); rpr.append(color)
    underline=OxmlElement("w:u"); underline.set(qn("w:val"),"single"); rpr.append(underline); run.append(rpr)
    t=OxmlElement("w:t"); t.text=text; run.append(t); hyperlink.append(run); paragraph._p.append(hyperlink)


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def add_figure(doc, filename, caption, alt):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(4)
    run=p.add_run(); shape=run.add_picture(str(FIGURES/filename), width=Inches(6.3))
    shape._inline.docPr.set("descr",alt)
    cap=doc.add_paragraph(); cap.style="Caption"; cap.alignment=WD_ALIGN_PARAGRAPH.CENTER; cap.paragraph_format.space_after=Pt(10)
    r=cap.add_run(caption); set_run(r,size=8.5,italic=True,color=MUTED)


def add_callout(doc, label_text, body, color=GREEN):
    table=doc.add_table(rows=1,cols=1); set_table_geometry(table,[9360]); cell=table.cell(0,0); set_cell_shading(cell,PALE)
    p=cell.paragraphs[0]; p.paragraph_format.space_after=Pt(0)
    r=p.add_run(label_text.upper()+"  "); set_run(r,size=9,bold=True,color=color)
    r=p.add_run(body); set_run(r,size=9.5,color=INK)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)


def configure_doc(doc: Document):
    section=doc.sections[0]
    section.top_margin=section.bottom_margin=section.left_margin=section.right_margin=Inches(TOKENS["page"]["margin"])
    section.header_distance=Inches(TOKENS["page"]["header"]); section.footer_distance=Inches(TOKENS["page"]["footer"])
    styles=doc.styles
    normal=styles["Normal"]; normal.font.name="Arial"; normal._element.rPr.rFonts.set(qn("w:ascii"),"Arial"); normal._element.rPr.rFonts.set(qn("w:hAnsi"),"Arial"); normal.font.size=Pt(TOKENS["body"]["size"]); normal.font.color.rgb=RGBColor.from_string(INK)
    normal.paragraph_format.space_after=Pt(TOKENS["body"]["after"]); normal.paragraph_format.line_spacing=TOKENS["body"]["line"]
    for name,key in (("Heading 1","h1"),("Heading 2","h2"),("Heading 3","h3")):
        st=styles[name]; token=TOKENS[key]; st.font.name=token["font"]; st._element.rPr.rFonts.set(qn("w:ascii"),token["font"]); st._element.rPr.rFonts.set(qn("w:hAnsi"),token["font"]); st.font.size=Pt(token["size"]); st.font.bold=True; st.font.color.rgb=RGBColor.from_string(token["color"]); st.paragraph_format.space_before=Pt(token["before"]); st.paragraph_format.space_after=Pt(token["after"]); st.paragraph_format.keep_with_next=True
    for name in ("List Bullet","List Number"):
        st=styles[name]; st.font.name="Arial"; st.font.size=Pt(10.5); st.paragraph_format.space_after=Pt(4); st.paragraph_format.line_spacing=1.18
    styles["Caption"].font.name="Arial"; styles["Caption"].font.size=Pt(8.5); styles["Caption"].font.color.rgb=RGBColor.from_string(MUTED)
    header=section.header; hp=header.paragraphs[0]; hp.text="WITNESS  /  TECHNICAL MONOGRAPH"; hp.alignment=WD_ALIGN_PARAGRAPH.LEFT
    for r in hp.runs: set_run(r,size=8,bold=True,color=MUTED)
    footer=section.footer; fp=footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    r=fp.add_run("AARYAN PATEL  /  AUGUST 2026  /  "); set_run(r,size=8,color=MUTED); add_page_field(fp)


def add_cover(doc: Document):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(95); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("ENGINEERING MONOGRAPH  /  WM-01"); set_run(r,size=10,bold=True,color=GREEN)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(22); p.paragraph_format.space_after=Pt(8)
    r=p.add_run("WITNESS"); set_run(r,size=34,bold=True,color=INK)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(20)
    r=p.add_run("Executable Evidence for Failure-Oriented Engineering Claims"); set_run(r,size=16,bold=True,color=CYAN)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("A receipt-gated reproducibility architecture across systems, ML infrastructure, autonomy, quantitative simulation, search, imaging, navigation, and numerical methods"); set_run(r,size=11.5,color=MUTED)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(34)
    r=p.add_run("AARYAN PATEL"); set_run(r,size=12,bold=True,color=INK)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("Independent Technical Monograph  |  August 2026  |  Version 1.0"); set_run(r,size=9.5,color=MUTED)
    add_figure(doc,"fig01-evidence-chain.png","Figure 1. WITNESS accepts a claim only when identity and behavior agree.","Evidence chain from claim and attack through oracle, source, receipt, execution, and accepted verdict.")
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("Not a university dissertation or peer-reviewed publication."); set_run(r,size=9,italic=True,color=AMBER)


def add_section(doc, title, paragraphs, level=1):
    doc.add_heading(title,level=level)
    for text in paragraphs:
        if text.startswith("I(") or text.startswith("B(") or text.startswith("The chamber verdict"):
            p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(.25); p.paragraph_format.right_indent=Inches(.25); p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(8)
            r=p.add_run(text); set_run(r,size=8.8,font="Menlo",color=INK)
        else:
            doc.add_paragraph(text)


def add_capsule_chapter(doc, number, c):
    doc.add_heading(f"{number}. {c['label']}: {c['system']}",level=1)
    add_callout(doc,f"{c['case']} / {c['runtime']}",c["claim"])
    table=doc.add_table(rows=1,cols=2)
    header = table.rows[0].cells
    header[0].text = "Field"
    header[1].text = "Pinned evidence"
    for cell in header:
        set_cell_shading(cell, INK)
        for run in cell.paragraphs[0].runs:
            set_run(run, size=8, bold=True, color="FFFFFF")
    rows=[("Repository",c["repo"]),("Pinned commit",c["commit"]),("Receipt SHA-256",c["receipt"]),("Adversarial scenario",c["attack"]),("Oracle",c["oracle"]),("Browser observation",c["observed"])]
    for idx,(a,b) in enumerate(rows):
        cells=table.add_row().cells; cells[0].text=a; cells[1].text=b
        set_cell_shading(cells[0],PALE if idx%2==0 else LIGHT); set_cell_shading(cells[1],"FFFFFF" if idx%2==0 else LIGHT)
        for r in cells[0].paragraphs[0].runs: set_run(r,size=8.5,bold=True,color=GREEN)
        for r in cells[1].paragraphs[0].runs: set_run(r,size=8.5,color=INK,font="Menlo" if idx in (1,2) else "Arial")
    set_table_geometry(table,[2050,7310])
    doc.add_paragraph()
    doc.add_heading("Mechanism and interpretation",level=2)
    for p in c["analysis"]: doc.add_paragraph(p)
    add_figure(doc,c["figure"],f"Figure {number-2}. {c['case']} - {c['system']} counterexample structure.",f"Technical diagram for {c['label']} capsule {c['case']}.")
    doc.add_heading("Claim boundary",level=2); doc.add_paragraph(c["boundary"])
    doc.add_heading("Reproduction target",level=2)
    p=doc.add_paragraph(); r=p.add_run(f"https://asp53826.github.io/witness/?capsule={c['id']}"); set_run(r,size=9,color=CYAN,font="Menlo")


def add_toc(doc):
    doc.add_heading("Contents",level=1)
    entries=[title for title,_ in FRONT_MATTER+CORE_SECTIONS]+[f"{i+6}. {c['label']}: {c['system']}" for i,c in enumerate(CAPSULES)]+[title for title,_ in SYNTHESIS]+["Appendix A. Reproduction Protocol","Appendix B. Capsule Identity Register","References"]
    for e in entries:
        p=doc.add_paragraph(style="List Bullet"); p.add_run(e)
    doc.add_heading("Artifact coordinates", level=2)
    coordinates = doc.add_table(rows=0, cols=2)
    for label_text, value in (
        ("Live chamber", "asp53826.github.io/witness"),
        ("Source", "github.com/asp53826/witness"),
        ("Publication", "PDF, Word, and GitHub Markdown editions"),
    ):
        cells = coordinates.add_row().cells
        cells[0].text = label_text
        cells[1].text = value
        set_cell_shading(cells[0], PALE)
        for run in cells[0].paragraphs[0].runs:
            set_run(run, size=8.5, bold=True, color=GREEN)
        for run in cells[1].paragraphs[0].runs:
            set_run(run, size=8.5, color=INK)
    set_table_geometry(coordinates, [2050, 7310])
    doc.add_page_break()


def add_back_cover(doc):
    doc.add_page_break()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(145)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("WITNESS")
    set_run(r, size=30, bold=True, color=INK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CLAIM  /  ATTACK  /  ORACLE  /  SOURCE  /  RECEIPT  /  RUN")
    set_run(r, size=10, bold=True, color=GREEN)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(26)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Executable evidence is strongest when the losing regime is easy to inspect and the claim boundary is impossible to miss.")
    set_run(r, size=13, color=MUTED)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(28)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_hyperlink(p, "RUN THE EIGHT CAPSULES", "https://asp53826.github.io/witness/")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_hyperlink(p, "INSPECT THE SOURCE", "https://github.com/asp53826/witness")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(42)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Independent technical monograph  /  Version 1.0  /  August 2026")
    set_run(r, size=9, color=MUTED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("No degree, peer-review status, or independent reproduction is claimed.")
    set_run(r, size=8.5, italic=True, color=AMBER)


def build_docx():
    doc=Document()
    properties = doc.core_properties
    properties.title = "WITNESS: Executable Evidence for Failure-Oriented Engineering Claims"
    properties.subject = "Independent technical monograph for the WITNESS reproducibility chamber"
    properties.author = "Aaryan Patel"
    properties.last_modified_by = "Aaryan Patel"
    properties.keywords = "reproducibility, counterexamples, engineering evidence, WebAssembly, Pyodide"
    properties.comments = "Version 1.0 — August 2026"
    configure_doc(doc); add_cover(doc)
    for title,paras in FRONT_MATTER: add_section(doc,title,paras)
    add_callout(doc,"Reading rule","A passing capsule means the pinned oracle correctly classified its intended regime. It does not mean the underlying method is universally correct.",AMBER)
    add_toc(doc)
    add_figure(doc,"fig02-runtime-topology.png","Figure 2. Runtime isolation and the shared structured result contract.","Browser topology with integrity inspector, WebAssembly, TypeScript, Pyodide Worker, and result receipt.")
    add_figure(doc,"fig03-capsule-matrix.png","Figure 3. Eight domains share one evidence protocol.","Eight WITNESS capsules arranged as a two-bank matrix.")
    for title,paras in CORE_SECTIONS: add_section(doc,title,paras)
    for i,c in enumerate(CAPSULES,start=6): add_capsule_chapter(doc,i,c)
    for title,paras in SYNTHESIS: add_section(doc,title,paras)
    doc.add_heading("Appendix A. Reproduction Protocol",level=1)
    steps=["Open a capsule-specific deep link.","Verify that format, receipt digest, embedded capsule digest, and commit binding show VERIFIED.","Execute the capsule without inverting the expected verdict.","Record measured metrics and trace lines.","Download the local run receipt and confirm signatureStatus is local-unattested.","Repeat with INVERT EXPECTED VERDICT and confirm that identical engine output yields a failing comparison.","For native benchmark claims, leave the browser and run the repository command at the pinned source revision."]
    for s in steps: doc.add_paragraph(s,style="List Number")
    doc.add_heading("Appendix B. Capsule Identity Register",level=1)
    table=doc.add_table(rows=1,cols=4); hdr=table.rows[0].cells
    for i,t in enumerate(("Case","Repository","Commit","Receipt SHA-256")): hdr[i].text=t; set_cell_shading(hdr[i],INK); [set_run(r,size=8,bold=True,color="FFFFFF") for r in hdr[i].paragraphs[0].runs]
    for c in CAPSULES:
        cells=table.add_row().cells
        vals=(c["case"],c["repo"],c["commit"][:12],c["receipt"][:16]+"...")
        for i,v in enumerate(vals): cells[i].text=v; [set_run(r,size=7.5,color=INK,font="Menlo" if i>1 else "Arial") for r in cells[i].paragraphs[0].runs]
    set_table_geometry(table,[900,2400,1900,4160])
    doc.add_heading("References",level=1)
    for index, ref in enumerate(REFERENCES, start=1):
        doc.add_paragraph(f"{index}.  {ref}")
    add_back_cover(doc)
    doc.save(DOCX_PATH)


def md_table(rows):
    out=["| Field | Value |","|---|---|"]
    for a,b in rows: out.append(f"| {a} | {str(b).replace('|','/')} |")
    return "\n".join(out)


def build_markdown():
    out=["# WITNESS", "", "## Executable Evidence for Failure-Oriented Engineering Claims", "", "**Independent Technical Monograph · Version 1.0 · August 2026**", "", "> **Status.** This is an independent engineering monograph, not a university dissertation, peer-reviewed publication, or claim that an academic degree has been awarded. No ACM badge or independent reproduction status is claimed.", "", "[Run WITNESS](https://asp53826.github.io/witness/) · [Download the designed PDF](../public/publications/WITNESS-Technical-Monograph.pdf) · [Download the Word edition](WITNESS-Technical-Monograph.docx)", "", "![WITNESS evidence chain](figures/fig01-evidence-chain.png)", ""]
    for title,paras in FRONT_MATTER:
        out += [f"## {title}",""] + sum(([p,""] for p in paras),[])
    out += ["## Contents","",*([f"- {t}" for t,_ in CORE_SECTIONS]+[f"- {i+6}. {c['label']}: {c['system']}" for i,c in enumerate(CAPSULES)]+[f"- {t}" for t,_ in SYNTHESIS]+["- Appendix A. Reproduction Protocol","- Appendix B. Capsule Identity Register","- References"]),""]
    out += ["![Browser execution topology](figures/fig02-runtime-topology.png)","","![Eight executable capsules](figures/fig03-capsule-matrix.png)",""]
    for title,paras in CORE_SECTIONS:
        out += [f"## {title}",""]
        for p in paras:
            out += (["```text",p,"```",""] if p.startswith(("I(","B(","The chamber verdict")) else [p,""])
    for i,c in enumerate(CAPSULES,start=6):
        out += [f"## {i}. {c['label']}: {c['system']}","",f"> **{c['case']} · {c['runtime']}**  ",f"> {c['claim']}",""]
        out += [md_table([("Repository",c["repo"]),("Pinned commit",f"`{c['commit']}`"),("Receipt SHA-256",f"`{c['receipt']}`"),("Adversarial scenario",c["attack"]),("Oracle",c["oracle"]),("Browser observation",c["observed"])]),""]
        out += ["### Mechanism and interpretation",""]
        for p in c["analysis"]: out += [p,""]
        out += [f"![{c['case']} technical diagram](figures/{c['figure']})","",f"### Claim boundary","",c["boundary"],"",f"**Reproduce:** https://asp53826.github.io/witness/?capsule={c['id']}",""]
    for title,paras in SYNTHESIS:
        out += [f"## {title}",""]
        for p in paras: out += [p,""]
    out += ["## Appendix A. Reproduction Protocol",""]
    steps=["Open a capsule-specific deep link.","Verify format, receipt digest, capsule digest, and commit binding.","Execute without inverting the expected verdict.","Record metrics and trace lines.","Download the local run receipt and confirm it is marked local-unattested.","Invert the expected verdict and confirm identical engine output yields a failed comparison.","Run the repository-native benchmark separately for any performance claim."]
    out += [f"{i}. {s}" for i,s in enumerate(steps,start=1)]+[""]
    out += ["## Appendix B. Capsule Identity Register","", "| Case | Repository | Commit | Receipt SHA-256 |","|---|---|---|---|"]
    for c in CAPSULES: out.append(f"| {c['case']} | {c['repo']} | `{c['commit']}` | `{c['receipt']}` |")
    out += ["","## References",""]+[f"{i}. {r}" for i,r in enumerate(REFERENCES,start=1)]+[""]
    MARKDOWN_PATH.write_text("\n".join(out),encoding="utf-8")


def main():
    generate_figures(); build_markdown(); build_docx()
    print(f"wrote {MARKDOWN_PATH}")
    print(f"wrote {DOCX_PATH}")
    print(f"wrote {len(list(FIGURES.glob('fig*.png')))} figures")


if __name__ == "__main__": main()
